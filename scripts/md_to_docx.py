#!/usr/bin/env python
"""
Converts paper/MemoryLifeBench_manuscript.md into a real, editable .docx
(Word-native headings/tables/bold/italic/code, not an image or a PDF
pretending to be one) so the paper can be opened directly in Microsoft
Word and edited -- tables/figures added, formatting tweaked, etc.

Handles the specific subset of Markdown this manuscript actually uses:
headings (#, ##, ###), paragraphs with inline **bold**/*italic*/`code`,
pipe tables, unordered lists (- item), ordered lists (1. item), and
blockquote lines (> ...) for the reproduced prompts (rendered as an
indented, shaded block in a monospace-ish font).

    python scripts/md_to_docx.py paper/MemoryLifeBench_manuscript.md paper/MemoryLifeBench_manuscript.docx
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`]+?`)")


def add_inline_runs(paragraph, text):
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 1:
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(chunk)


def set_cell_shading(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, rows):
    header, *body = rows
    n_cols = len(header)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        set_cell_shading(cell, "D9E2F3")
    for row_vals in body:
        row = table.add_row()
        for i, text in enumerate(row_vals):
            if i >= n_cols:
                break
            cell = row.cells[i]
            cell.text = ""
            add_inline_runs(cell.paragraphs[0], text)
    doc.add_paragraph()


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def main():
    src, dst = sys.argv[1], sys.argv[2]
    lines = open(src, encoding="utf-8").read().splitlines()

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    first_heading_done = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Title (first line, single #)
        if stripped.startswith("# ") and not first_heading_done:
            p = doc.add_heading(level=0)
            add_inline_runs(p, stripped[2:].strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_heading_done = True
            i += 1
            continue

        if stripped == "Bhargav Shendge":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.italic = True
            run.font.size = Pt(12)
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(level=2)
            add_inline_runs(p, stripped[4:].strip())
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_heading(level=1)
            add_inline_runs(p, stripped[3:].strip())
            i += 1
            continue

        # Table: header line + separator line
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [parse_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # Blockquote block (prompts)
        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and (lines[i].strip().startswith(">") or lines[i].strip() == ""):
                if lines[i].strip().startswith(">"):
                    text = lines[i].strip()[1:].strip()
                    quote_lines.append(text)
                elif quote_lines and quote_lines[-1] != "":
                    quote_lines.append("")
                i += 1
                if i < len(lines) and not lines[i].strip().startswith(">") and lines[i].strip() != "":
                    break
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            for j, ql in enumerate(quote_lines):
                if j > 0:
                    p.add_run().add_break()
                run = p.add_run(ql if ql else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Unordered list
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:].strip())
            i += 1
            continue

        # Ordered list
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m.group(2))
            i += 1
            continue

        # Regular paragraph (collect until blank line, table, heading, quote, or list)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            nxt_stripped = nxt.strip()
            if not nxt_stripped:
                break
            if nxt_stripped.startswith(("#", "|", ">", "- ")) or re.match(r"^\d+\.\s", nxt_stripped):
                break
            para_lines.append(nxt_stripped)
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(para_lines))
        p.paragraph_format.space_after = Pt(8)

    doc.save(dst)
    print(f"written -> {dst}")


if __name__ == "__main__":
    main()
